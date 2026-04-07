# -*- coding: utf-8 -*-
"""
Generateur DOCX — Rapport U62 V2
Style inspire de RAPPORT_U62_WOW.docx
BAHAFID Mohamed — BTS MEC Session 2026
"""
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from contenu_v2 import (
    CANDIDAT, PAGE_01, PAGE_02, PAGE_03, PAGE_04, PAGE_06, PAGE_07,
    PAGE_09, PAGE_10, PAGE_15, PAGE_17, PAGE_18, PAGE_20, PAGE_21,
    PAGE_23, PAGE_24, PAGE_25, PAGE_26, PAGE_27, PAGE_28,
    SITUATION_1, SITUATION_2, SITUATION_3, SITUATION_4, SITUATION_5,
    TABLE_CORPS_ETAT, TABLE_BUDGET_COMMUNES, TABLE_BUDGET_ROUTE,
    TABLE_AUTRES_MARCHES, TABLE_COMPARAISON_REG,
)

BASE = r"D:\PREPA BTS MEC\08_U62_Rapport_Activites"
OUT  = os.path.join(BASE, "Rapport_Redaction", "RAPPORT_U62_V2.docx")

# ─── Palette couleurs ────────────────────────────────────────────────────────
CNV = "1E3A5F"   # Navy (corps de page)
COG = "F39200"   # Orange
CTQ = "5CC8C0"   # Teal
CWH = "FFFFFF"   # Blanc
CLG = "F0F4F8"   # Gris clair
CDK = "2D2D2D"   # Très sombre
CGY = "888888"   # Gris
CLO = "FFF8F0"   # Orange très clair
CDN = "0F1B2D"   # Navy très foncé (couverture)
CBL = "2E86C1"   # Bleu WOW
COR = "CA6F1E"   # Orange WOW
CTL = "148F77"   # Teal WOW

# ─── Helpers XML ────────────────────────────────────────────────────────────
def _qn(tag): return qn(tag)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(_qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(_qn('w:val'), 'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, sides, color, sz=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = tcPr.find(_qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders')
        tcPr.append(tcB)
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(_qn('w:val'), 'single')
        b.set(_qn('w:sz'), str(sz))
        b.set(_qn('w:color'), color)
        b.set(_qn('w:space'), '0')
        tcB.append(b)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = tcPr.find(_qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders')
        tcPr.append(tcB)
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(_qn('w:val'), 'none')
        tcB.append(b)

def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(_qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(_qn('w:val'), 'none')
        tblB.append(b)
    tblPr.append(tblB)

def set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    pSpacing = pPr.find(_qn('w:spacing'))
    if pSpacing is None:
        pSpacing = OxmlElement('w:spacing')
        pPr.append(pSpacing)
    pSpacing.set(_qn('w:before'), str(int(before * 20)))
    pSpacing.set(_qn('w:after'), str(int(after * 20)))
    pSpacing.set(_qn('w:line'), '276')
    pSpacing.set(_qn('w:lineRule'), 'auto')

def set_cell_vmid(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(_qn('w:val'), 'center')
    tcPr.append(vAlign)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(_qn('w:w'), str(int(width_cm * 567)))  # 567 = twips per cm
    tcW.set(_qn('w:type'), 'dxa')
    tcPr.append(tcW)

# ─── Helpers mise en forme texte ─────────────────────────────────────────────
def styled_run(para, text, bold=False, italic=False, size_pt=10.5,
               color=CDK, font='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def heading_para(doc, text, level=1, color=CNV, size_pt=14, font='Calibri'):
    """Titre sans style Word — formatage manuel."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, text, bold=True, size_pt=size_pt, color=color, font=font)
    return p

def body_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              color=CDK, size_pt=10.5, before=0, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    styled_run(p, text, size_pt=size_pt, color=color)
    return p

def bullet_para(doc, text, color=CDK, size_pt=10, icon='▸', icon_color=COG):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, f'{icon}  ', bold=True, size_pt=size_pt, color=icon_color)
    styled_run(p, text, size_pt=size_pt, color=color)
    return p

# ─── Blocs structurels ───────────────────────────────────────────────────────
def section_banner(doc, title, subtitle='', bg=CNV, fg=CWH):
    """Bandeau coloré pleine largeur pour titre de section."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders_none(tbl)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.width = Cm(16.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    styled_run(p, title, bold=True, size_pt=16, color=fg, font='Calibri')
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(6)
        styled_run(p2, subtitle, size_pt=9, color='CCE5F5', font='Calibri')
    doc.add_paragraph()
    return tbl

def mini_banner(doc, text, bg=COG, fg=CWH):
    """Petit bandeau pour sous-titre."""
    tbl = doc.add_table(rows=1, cols=1)
    set_table_borders_none(tbl)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, text, bold=True, size_pt=10, color=fg, font='Calibri')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def kv_table(doc, rows_data, col_widths=(5.5, 11)):
    """Tableau clé→valeur (2 colonnes)."""
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    set_table_borders_none(tbl)
    for i, (k, v) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])
        bg = CLG if i % 2 == 0 else CWH
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        pk = row.cells[0].paragraphs[0]
        pk.paragraph_format.space_before = Pt(4)
        pk.paragraph_format.space_after = Pt(4)
        styled_run(pk, k, bold=True, size_pt=9.5, color=CNV)
        pv = row.cells[1].paragraphs[0]
        pv.paragraph_format.space_before = Pt(4)
        pv.paragraph_format.space_after = Pt(4)
        styled_run(pv, v, size_pt=9.5, color=CDK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def data_table(doc, headers, rows, col_widths=None, header_bg=CNV):
    """Tableau de données avec en-tête coloré."""
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    set_table_borders_none(tbl)
    # En-tête
    hrow = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        set_cell_bg(cell, header_bg)
        if col_widths:
            cell.width = Cm(col_widths[j])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        styled_run(p, h, bold=True, size_pt=8.5, color=CWH)
    # Lignes
    for i, row_data in enumerate(rows):
        drow = tbl.rows[i + 1]
        is_total = str(row_data[0]).upper().startswith('TOTAL')
        bg = CLO if is_total else (CLG if i % 2 == 0 else CWH)
        for j, val in enumerate(row_data):
            cell = drow.cells[j]
            set_cell_bg(cell, bg)
            if col_widths:
                cell.width = Cm(col_widths[j])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            text_color = COG if is_total else CDK
            styled_run(p, str(val), bold=is_total, size_pt=9, color=text_color)
        # Fine ligne séparatrice
        for cell in drow.cells:
            set_cell_borders(cell, ['bottom'], 'DDDDDD', sz=4)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def cpar_box(doc, situation):
    """Bloc CPAR complet pour une situation professionnelle."""
    # Bandeau situation
    tbl_h = doc.add_table(rows=1, cols=2)
    set_table_borders_none(tbl_h)
    hdr = tbl_h.rows[0]
    num_cell = hdr.cells[0]
    set_cell_bg(num_cell, COG)
    num_cell.width = Cm(3)
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    styled_run(p, f'SITUATION {situation["numero"]}', bold=True, size_pt=9, color=CWH)
    title_cell = hdr.cells[1]
    set_cell_bg(title_cell, CNV)
    title_cell.width = Cm(13.5)
    pt = title_cell.paragraphs[0]
    pt.paragraph_format.space_before = Pt(8)
    pt.paragraph_format.space_after = Pt(4)
    styled_run(pt, situation['titre'], bold=True, size_pt=11, color=CWH)
    p2 = title_cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    styled_run(p2, f"Compétence : {situation['competence']}", size_pt=8.5, color='AAD4F0')
    # Indicateur clé
    kpi_tbl = doc.add_table(rows=1, cols=2)
    set_table_borders_none(kpi_tbl)
    kc = kpi_tbl.rows[0].cells[0]
    set_cell_bg(kc, CLO)
    kc.width = Cm(4)
    set_cell_borders(kc, ['left'], COG, sz=16)
    pk = kc.paragraphs[0]
    pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pk.paragraph_format.space_before = Pt(8)
    pk.paragraph_format.space_after = Pt(8)
    styled_run(pk, situation['chiffre_cle'], bold=True, size_pt=20, color=COG)
    p2k = kc.add_paragraph()
    p2k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2k.paragraph_format.space_before = Pt(0)
    p2k.paragraph_format.space_after = Pt(8)
    styled_run(p2k, situation['chiffre_label'], size_pt=8, color=CDK)
    # Sections CPAR
    sections = [
        ('C', 'CONTEXTE', situation['contexte'], CBL),
        ('P', 'PROBLÈME', situation['probleme'], 'C0392B'),
        ('A', 'ACTION', situation['action'], CTL),
        ('R', 'RÉSULTAT', situation['resultat'], COR),
    ]
    txt_cell = kpi_tbl.rows[0].cells[1]
    txt_cell.width = Cm(12.5)
    set_cell_bg(txt_cell, CWH)
    for letter, label, content, lcolor in sections:
        ptxt = txt_cell.add_paragraph()
        ptxt.paragraph_format.space_before = Pt(4)
        ptxt.paragraph_format.space_after = Pt(2)
        styled_run(ptxt, f'{letter} — {label}  ', bold=True, size_pt=9, color=lcolor)
        styled_run(ptxt, content, size_pt=9, color=CDK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def phase_timeline(doc, phases):
    """Affiche les 5 phases en tableau."""
    tbl = doc.add_table(rows=len(phases), cols=3)
    set_table_borders_none(tbl)
    colors = [CTQ, CNV, COG, CTQ, COG]
    for i, (date, title, desc) in enumerate(phases):
        row = tbl.rows[i]
        # Date
        dc = row.cells[0]
        set_cell_bg(dc, colors[i])
        dc.width = Cm(3)
        pd = dc.paragraphs[0]
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pd.paragraph_format.space_before = Pt(6)
        pd.paragraph_format.space_after = Pt(6)
        styled_run(pd, date, bold=True, size_pt=8.5, color=CWH)
        # Titre
        tc2 = row.cells[1]
        set_cell_bg(tc2, CLG)
        tc2.width = Cm(4.5)
        pt2 = tc2.paragraphs[0]
        pt2.paragraph_format.space_before = Pt(6)
        pt2.paragraph_format.space_after = Pt(2)
        styled_run(pt2, title, bold=True, size_pt=9.5, color=CNV)
        # Description
        dc3 = row.cells[2]
        set_cell_bg(dc3, CWH)
        dc3.width = Cm(9)
        pd3 = dc3.paragraphs[0]
        pd3.paragraph_format.space_before = Pt(6)
        pd3.paragraph_format.space_after = Pt(6)
        styled_run(pd3, desc, size_pt=9, color=CDK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def defi_table(doc, defis):
    """Tableau défis/problème/solution (3 colonnes)."""
    headers = ['Défi', 'Problème rencontré', 'Solution apportée']
    rows = [(d, p, s) for d, p, s in defis]
    data_table(doc, headers, rows, col_widths=[3.5, 6.5, 6.5], header_bg=COG)

def kpi_band(doc, kpis):
    """Bande de KPIs."""
    tbl = doc.add_table(rows=1, cols=len(kpis))
    set_table_borders_none(tbl)
    kpi_colors = [COG, CNV, CTQ, 'C0392B', COG]
    for i, (v, l) in enumerate(kpis):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, kpi_colors[i % len(kpi_colors)])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        styled_run(p, v, bold=True, size_pt=16, color=CWH)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        styled_run(p2, l, size_pt=8, color='FFEECC')
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def reflexif_bloc(doc, title, text, letter, bg, border_color):
    """Bloc bilan réflexif avec bordure gauche colorée."""
    tbl = doc.add_table(rows=1, cols=2)
    set_table_borders_none(tbl)
    lc = tbl.rows[0].cells[0]
    set_cell_bg(lc, border_color)
    lc.width = Cm(1)
    pl = lc.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.space_before = Pt(10)
    pl.paragraph_format.space_after = Pt(10)
    styled_run(pl, letter, bold=True, size_pt=14, color=CWH)
    rc = tbl.rows[0].cells[1]
    set_cell_bg(rc, bg)
    rc.width = Cm(15.5)
    pt = rc.paragraphs[0]
    pt.paragraph_format.space_before = Pt(8)
    pt.paragraph_format.space_after = Pt(4)
    styled_run(pt, title, bold=True, size_pt=11, color=border_color)
    pb = rc.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pb.paragraph_format.space_before = Pt(4)
    pb.paragraph_format.space_after = Pt(10)
    styled_run(pb, text, size_pt=9.5, color=CDK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def horizon_card(doc, horizons):
    """Tableau 3 colonnes pour les horizons pro."""
    tbl = doc.add_table(rows=2, cols=3)
    set_table_borders_none(tbl)
    hz_colors = [COG, CTQ, CNV]
    for i, (label, date, desc) in enumerate(horizons):
        hc = tbl.rows[0].cells[i]
        set_cell_bg(hc, hz_colors[i])
        hc.width = Cm(5.5)
        ph = hc.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(6)
        ph.paragraph_format.space_after = Pt(2)
        styled_run(ph, label, bold=True, size_pt=10, color=CWH)
        pd2 = hc.add_paragraph()
        pd2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pd2.paragraph_format.space_before = Pt(0)
        pd2.paragraph_format.space_after = Pt(6)
        styled_run(pd2, date, size_pt=9, color='FFEECC')
        # Desc
        dc = tbl.rows[1].cells[i]
        set_cell_bg(dc, CLG)
        dc.width = Cm(5.5)
        desc_clean = desc.replace('\n· ', '\n▸ ').replace('\n', '\n')
        for line in desc_clean.strip().split('\n'):
            pl = dc.add_paragraph()
            pl.paragraph_format.space_before = Pt(2)
            pl.paragraph_format.space_after = Pt(2)
            if line.strip().startswith('▸'):
                pl.paragraph_format.left_indent = Cm(0.3)
                styled_run(pl, '▸ ', bold=True, size_pt=9, color=hz_colors[i])
                styled_run(pl, line.strip()[1:].strip(), size_pt=9, color=CDK)
            else:
                styled_run(pl, line.strip(), bold=line.strip().startswith('Axe') or line.strip().startswith('BIMCO'),
                           size_pt=9, color=CNV if line.strip().startswith('Axe') else CDK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# ════════════════════════════════════════════════════════════════════════════
#  CONSTRUCTION DU DOCUMENT
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# ─── Mise en page A4 ────────────────────────────────────────────────────────
sect = doc.sections[0]
sect.page_width  = Cm(21)
sect.page_height = Cm(29.7)
sect.left_margin   = Cm(2.5)
sect.right_margin  = Cm(2)
sect.top_margin    = Cm(2)
sect.bottom_margin = Cm(2)

# Style Normal par défaut
from docx.oxml import OxmlElement as OE
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE COUVERTURE
# ══════════════════════════════════════════════════════════════════════════════
# Bandeau top
tbl_top = doc.add_table(rows=1, cols=1)
set_table_borders_none(tbl_top)
top_c = tbl_top.rows[0].cells[0]
set_cell_bg(top_c, CDN)
top_c.width = Cm(16.5)
pt = top_c.paragraphs[0]
pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt.paragraph_format.space_before = Pt(10)
pt.paragraph_format.space_after = Pt(4)
styled_run(pt, 'BTS MANAGEMENT ÉCONOMIQUE DE LA CONSTRUCTION — SESSION 2026', bold=True, size_pt=8, color='AABBCC')
pt2 = top_c.add_paragraph()
pt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pt2.paragraph_format.space_before = Pt(0)
pt2.paragraph_format.space_after = Pt(10)
styled_run(pt2, 'Académie de Lyon  |  Candidat n° 02537399911', size_pt=8, color='8899AA')

# Titre principal
tbl_main = doc.add_table(rows=1, cols=1)
set_table_borders_none(tbl_main)
mc = tbl_main.rows[0].cells[0]
set_cell_bg(mc, CNV)
mc.width = Cm(16.5)
pm = mc.paragraphs[0]
pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm.paragraph_format.space_before = Pt(30)
pm.paragraph_format.space_after = Pt(6)
styled_run(pm, 'RAPPORT', bold=True, size_pt=40, color=COG)
pm2 = mc.add_paragraph()
pm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm2.paragraph_format.space_before = Pt(0)
pm2.paragraph_format.space_after = Pt(6)
styled_run(pm2, "D'ACTIVITÉS", bold=True, size_pt=40, color=CWH)
pm3 = mc.add_paragraph()
pm3.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm3.paragraph_format.space_before = Pt(0)
pm3.paragraph_format.space_after = Pt(16)
styled_run(pm3, "PROFESSIONNELLES", bold=True, size_pt=40, color=CWH)
# Ligne orange
pm4 = mc.add_paragraph()
pm4.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm4.paragraph_format.space_before = Pt(4)
pm4.paragraph_format.space_after = Pt(16)
styled_run(pm4, '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━', bold=False, size_pt=12, color=COG)
# Sous-titre
pm5 = mc.add_paragraph()
pm5.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm5.paragraph_format.space_before = Pt(0)
pm5.paragraph_format.space_after = Pt(30)
styled_run(pm5, 'BAHAFID Mohamed  —  BTS MEC  —  Session 2026', bold=True, size_pt=14, color='CCE5FF')

# Bandeau bas couverture
tbl_bot = doc.add_table(rows=1, cols=3)
set_table_borders_none(tbl_bot)
bot_data = [
    ("Candidat",  "BAHAFID Mohamed\nn° 02537399911"),
    ("Structure", "Conseil Régional de\nBéni Mellal-Khénifra"),
    ("Activité",  "BIMCO — Projeteur BIM\nÉconomiste de la construction"),
]
bot_colors = [CDN, CNV, CDN]
for i, (label, val) in enumerate(bot_data):
    c = tbl_bot.rows[0].cells[i]
    set_cell_bg(c, bot_colors[i])
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, label.upper(), bold=True, size_pt=7, color='7799BB')
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    styled_run(p2, val.replace('\n', '\n'), bold=True, size_pt=9, color=CWH)

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FICHE D'IDENTITÉ
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, 'FICHE D\'IDENTITÉ DU CANDIDAT', '', bg=CNV)
kv_table(doc, PAGE_02["champs"], col_widths=[5.5, 11])
body_para(doc, PAGE_02["pied"], align=WD_ALIGN_PARAGRAPH.CENTER, color=CGY, size_pt=8)
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, 'SOMMAIRE', bg=CDN)
for (num, titre, desc, page) in PAGE_03["sections"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if num != "—":
        styled_run(p, f'  {num}. ', bold=True, size_pt=11, color=COG)
    else:
        styled_run(p, '  ● ', bold=True, size_pt=11, color=CTQ)
    styled_run(p, f'{titre}', bold=True, size_pt=11, color=CNV)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.2)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    styled_run(p2, desc + f'  ...... {page}', size_pt=9, color=CGY)
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION ET PARCOURS
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, 'INTRODUCTION ET PARCOURS', bg=CBL)
body_para(doc, PAGE_04["intro_texte"], before=0, after=8)
mini_banner(doc, 'MON PARCOURS EN 5 PHASES', bg=CNV)
phase_timeline(doc, PAGE_04["phases"])
mini_banner(doc, 'PROJETS ANALYSÉS', bg=COG)
for proj in PAGE_04["projets"]:
    bullet_para(doc, proj)
mini_banner(doc, 'CHIFFRES CLÉS', bg=CTQ)
kpi_band(doc, PAGE_04["chiffres_cles"])
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 01 — CADRE PROFESSIONNEL
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, '01 — CADRE PROFESSIONNEL',
               'Le Conseil Régional de Béni Mellal-Khénifra | BIMCO', bg=CDN)

# Conseil Régional
heading_para(doc, '1.1 Conseil Régional de Béni Mellal-Khénifra', size_pt=13, color=CBL)
body_para(doc,
    "Collectivité territoriale créée dans le cadre du découpage régional de 2015, le Conseil Régional de Béni Mellal-Khénifra "
    "couvre cinq provinces (Béni Mellal, Azilal, Fquih Ben Salah, Khénifra et Khouribga) sur 28 374 km² pour 2,5 millions d'habitants. "
    "L'Agence d'Exécution des Projets, dirigée par M. Abdelaaziz DOGHMANI, assure la maîtrise d'ouvrage de l'ensemble des "
    "infrastructures régionales : routes rurales, voirie urbaine, adduction d'eau potable et équipements collectifs. "
    "L'ensemble des marchés de travaux est régi par le Décret n°2-12-349 du 20 mars 2013 relatif aux marchés publics — "
    "appel d'offres ouvert au-dessus de 500 000 DH, pièces constitutives obligatoires : CPS, RC, BPDE, plans et estimation confidentielle.",
    before=0, after=6)
for f in PAGE_06["facts"]:
    bullet_para(doc, f'{f[0]} : {f[1]}')
mini_banner(doc, 'MISSIONS EN TANT QUE TECHNICIEN DE SUIVI', bg=CNV)
for m in PAGE_06["missions"]:
    bullet_para(doc, m)

# BIMCO
heading_para(doc, '1.2 BIMCO — Mon activité indépendante', size_pt=13, color=CBL)
body_para(doc,
    "BIMCO a été créé le 9 janvier 2026 (SIREN 999580053, APE 7112B — Ingénierie, études techniques) à Bussières (Loire). "
    "Le nom reflète le positionnement à la croisée du BIM (Building Information Modeling) et de l'économie de la COnstruction. "
    "BIMCO comble un vide réel : les outils BIM sont conçus pour les architectes et ingénieurs — très peu d'économistes "
    "de la construction maîtrisent le BIM et le développement d'outils numériques. "
    "BIMCO apporte la rigueur de l'économiste terrain avec les outils du BIM.",
    before=0, after=6)
mini_banner(doc, 'DOMAINES D\'INTERVENTION', bg=COG)
for d, desc in PAGE_07["domaines"]:
    bullet_para(doc, f'{d} : {desc}')
body_para(doc,
    f"Application phare : {PAGE_07['app']['titre']} — {PAGE_07['app']['url']} — Stack : {PAGE_07['app']['stack']}",
    color=CBL, size_pt=9)
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 02 — PROJET 1 : MISE À NIVEAU 4 COMMUNES
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, '02 — PROJET 1 : MISE À NIVEAU DE 4 COMMUNES',
               '53,5 M DH TTC — 8 corps d\'état — Province de Khénifra', bg=COG)

heading_para(doc, '2.1 Fiche d\'identité du projet', size_pt=13, color=CBL)
body_para(doc,
    "Le marché n°38-RBK-2017 (Lot 4) avait pour objectif la mise à niveau des centres de quatre communes rurales "
    "de la province de Khénifra. Ce programme d'aménagement urbain et VRD de 53,5 M DH TTC couvre huit corps d'état distincts "
    "— de l'assainissement à l'éclairage public — sur quatre sites géographiquement dispersés de 20 à 80 km.",
    before=0, after=6)
kv_table(doc, [(k, v.replace('\n', ', ')) for k, v in PAGE_09["fiche"]])

heading_para(doc, '2.2 Budget et répartition par commune', size_pt=12, color=CBL)
data_table(doc,
    TABLE_BUDGET_COMMUNES["colonnes"],
    TABLE_BUDGET_COMMUNES["lignes"],
    col_widths=[4.5, 4, 4, 4])

heading_para(doc, '2.3 Les 8 corps d\'état', size_pt=12, color=CBL)
data_table(doc,
    TABLE_CORPS_ETAT["colonnes"],
    TABLE_CORPS_ETAT["lignes"],
    col_widths=[1.5, 4.5, 10.5])

# Situations CPAR
mini_banner(doc, '5 SITUATIONS PROFESSIONNELLES — DÉMARCHE CPAR', bg=CDN)
body_para(doc,
    "Chaque situation est analysée selon la démarche CPAR : le Contexte pose la mission et les enjeux, "
    "le Problème identifie l'obstacle concret rencontré, l'Action décrit la réponse apportée avec les moyens disponibles, "
    "le Résultat mesure l'impact réel et les enseignements tirés.",
    before=0, after=8)
cpar_box(doc, SITUATION_1)
cpar_box(doc, SITUATION_2)
cpar_box(doc, SITUATION_3)
cpar_box(doc, SITUATION_4)

heading_para(doc, '2.4 Difficultés rencontrées et solutions apportées', size_pt=12, color=CBL)
defi_table(doc, PAGE_15["defis"])
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 03 — PROJET 2 : ROUTE LEHRI-KERROUCHEN
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, '03 — PROJET 2 : ROUTE LEHRI-KERROUCHEN',
               '29 M DH TTC — 25 km en zone montagneuse — Programme PRR3', bg=CTQ)

heading_para(doc, '3.1 Fiche d\'identité du projet', size_pt=13, color=CBL)
body_para(doc,
    "La construction de la route nationale communale (RNC) reliant Lehri et Kerrouchen s'inscrit dans le "
    "Programme National des Routes Rurales (PRR3), visant à désenclaver les zones rurales du Moyen Atlas. "
    "Le tracé de 25 km traverse un terrain accidenté avec un dénivelé cumulé de 400 m et des pentes atteignant 12% en lacets. "
    "Le marché n°46-RBK-2017 comprend 53 prix sur trois sections : linéaire principal (23 prix), carrefour (11 prix) et bretelles (19 prix).",
    before=0, after=6)
kv_table(doc, [(k, v) for k, v in PAGE_17["fiche"]])

heading_para(doc, '3.2 Principaux métrés', size_pt=12, color=CBL)
data_table(doc,
    ['Quantité', 'Désignation'],
    [(v, l) for v, l in PAGE_17["metres"]],
    col_widths=[3.5, 13])

heading_para(doc, '3.3 Répartition budgétaire', size_pt=12, color=CBL)
data_table(doc,
    TABLE_BUDGET_ROUTE["colonnes"],
    TABLE_BUDGET_ROUTE["lignes"],
    col_widths=[9.5, 4, 3])
body_para(doc, f"Note : {PAGE_18['callout'].replace(chr(10), ' — ')}", color=COR, size_pt=9)

# Situation 5
mini_banner(doc, 'SITUATION 5 — DÉMARCHE CPAR', bg=CDN)
cpar_box(doc, SITUATION_5)

heading_para(doc, '3.4 Défis d\'un chantier routier en zone montagneuse', size_pt=12, color=CBL)
defi_table(doc, PAGE_20["defis"])
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ACTIVITÉS COMPLÉMENTAIRES
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, 'ACTIVITÉS COMPLÉMENTAIRES',
               '5 autres marchés au Conseil Régional | Expérience terrain France', bg=CNV)

heading_para(doc, '4.1 Autres marchés suivis au Conseil Régional', size_pt=13, color=CBL)
body_para(doc,
    "En parallèle des deux projets principaux, j'ai participé au suivi de cinq autres marchés de travaux, "
    "couvrant des types d'infrastructures variés : routes rurales, pistes, adduction d'eau potable et voirie urbaine. "
    "Sur ces marchés : vérification des implantations, participation aux commissions d'ouverture des plis, "
    "suivi des ordres de service et vérification ponctuelle des situations de travaux.",
    before=0, after=6)
data_table(doc,
    ['N° Marché', 'Objet des travaux', 'Type / Montant'],
    TABLE_AUTRES_MARCHES["lignes"],
    col_widths=[2.5, 9, 5])

heading_para(doc, '4.2 Expérience terrain en France (2022-2024)', size_pt=13, color=CBL)
body_para(doc,
    "L'expérience terrain en France a profondément enrichi ma vision de l'économiste de la construction. "
    "Comprendre les coûts réels de production — main-d'œuvre, rendements, consommation de matériaux, coûts de matériel — "
    "rend les estimations infiniment plus justes. Maîtriser les techniques de gros œuvre (coffrage, ferraillage, bétonnage) "
    "est indispensable pour réaliser des métrés précis et analyser les offres avec pertinence.",
    before=0, after=6)
for d, p, dt in PAGE_21["france"]:
    bullet_para(doc, f'{d} — {p} : {dt}')
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 04 — BILAN ET ANALYSE
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, '04 — BILAN ET ANALYSE',
               'Compétences | Comparaison Maroc/France | Bilan réflexif | BIM | Projet pro', bg=CDN)

# Synthèse compétences
heading_para(doc, '5.1 Synthèse des activités et compétences', size_pt=13, color=CBL)
body_para(doc,
    "Ce tableau synthétise les activités professionnelles réalisées sur les deux projets et les cinq situations CPAR, "
    "croisées avec les compétences du BTS MEC mobilisées. La triple compétence MOA + Exécution + BIM "
    "permet d'établir des estimations réalistes, d'analyser les offres avec pertinence et d'exploiter les maquettes numériques "
    "pour des métrés précis. Niveau Maîtrise : pratique régulière et autonome. Niveau Expert : transmission aux collaborateurs.",
    before=0, after=6)
data_table(doc,
    ['Activité réalisée', 'Sous-compétence mobilisée', 'Situation', 'Niveau'],
    [(a, sc, sit, niv) for a, sc, sit, niv in PAGE_23["tableau"]],
    col_widths=[6.5, 4.5, 2.5, 3])

# Comparaison Maroc / France
heading_para(doc, '5.2 Analyse comparative Maroc / France', size_pt=13, color=CBL)
body_para(doc, PAGE_24["intro"], before=0, after=6)
data_table(doc,
    ['Aspect', 'Maroc', 'France'],
    TABLE_COMPARAISON_REG["lignes"],
    col_widths=[4, 6.5, 6])
body_para(doc, PAGE_24["synthese"], color=CNV, size_pt=9.5, before=4, after=8)

# Bilan réflexif
heading_para(doc, '5.3 Bilan réflexif', size_pt=13, color=CBL)
body_para(doc,
    "Avec le recul, chaque situation m'a appris quelque chose sur ma façon de travailler. "
    "Certains réflexes sont venus tardivement — le tableau de bord, la demande d'études complémentaires, "
    "la traçabilité photographique. Ce bilan tire les leçons de 8 années de pratique pour construire la suite.",
    before=0, after=6)
bloc_styles = [('A', CLO, COG), ('D', 'E8F8F6', CTQ), ('M', 'EDF1F5', CNV)]
for i, (t, tx) in enumerate(PAGE_25["blocs"]):
    letter, bg, bc = bloc_styles[i]
    reflexif_bloc(doc, t, tx, letter, bg, bc)

# Protocole BIM
heading_para(doc, '5.4 Protocole de collaboration BIM', size_pt=13, color=CBL)
body_para(doc,
    "Le BIM transforme la chaîne métré → estimation → chiffrage en la rendant plus fiable et plus rapide. "
    "Ma formation Technicien Modeleur BIM (AFPA Colmar, 8 mois) m'a permis de modéliser un bâtiment R+2 "
    "et d'en extraire automatiquement 78 postes de métrés — avec un écart de seulement 1,8% par rapport "
    "au métré manuel traditionnel. La détection de 12 clashs structure/réseaux illustre la valeur ajoutée concrète.",
    before=0, after=6)
mini_banner(doc, 'CONVENTION BIM APPLIQUÉE', bg=CTQ)
kv_table(doc, PAGE_26["convention"]["items"], col_widths=[4, 12.5])
mini_banner(doc, 'WORKFLOW BIM', bg=CNV)
for step in PAGE_26["workflow"]:
    bullet_para(doc, step, icon_color=CTQ)
mini_banner(doc, PAGE_26["cas_concret"]["titre"], bg=COG)
body_para(doc, PAGE_26["cas_concret"]["details"], before=0, after=4)
body_para(doc, PAGE_26["apport_mec"], color=CBL, size_pt=9, before=4, after=8)

# Projet professionnel
heading_para(doc, '5.5 Mon projet professionnel', size_pt=13, color=CBL)
body_para(doc,
    "Mon projet repose sur une conviction forte : les outils numériques doivent être au service de l'économiste "
    "de la construction — et non l'inverse. Le marché de l'ingénierie BIM est dominé par des architectes et ingénieurs ; "
    "très peu d'économistes maîtrisent à la fois le BIM, le développement d'outils et la réalité du terrain. "
    "BIMCO occupe ce créneau rare : la rigueur de l'économiste terrain combinée aux outils du BIM (Revit, Dynamo, plugins C#, apps web).",
    before=0, after=8)
horizon_card(doc, PAGE_27["horizons"])
tbl_cite = doc.add_table(rows=1, cols=1)
set_table_borders_none(tbl_cite)
cc = tbl_cite.rows[0].cells[0]
set_cell_bg(cc, CLO)
set_cell_borders(cc, ['left'], COG, sz=20)
pc = cc.paragraphs[0]
pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc.paragraph_format.space_before = Pt(12)
pc.paragraph_format.space_after = Pt(12)
styled_run(pc, PAGE_27["citation"].replace('\n', ' '), bold=True, size_pt=10, color=CNV)
page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
section_banner(doc, 'CONCLUSION', bg=CDN)
body_para(doc, PAGE_28["resume"], before=0, after=10)
mini_banner(doc, 'CHIFFRES CLÉS DU PARCOURS', bg=COG)
kpi_band(doc, PAGE_28["kpis"])
mini_banner(doc, 'ENSEIGNEMENTS PRINCIPAUX', bg=CNV)
for titre, texte in PAGE_28["points"]:
    bullet_para(doc, f'{titre} : {texte}')
doc.add_paragraph().paragraph_format.space_after = Pt(10)
tbl_final = doc.add_table(rows=1, cols=1)
set_table_borders_none(tbl_final)
fc = tbl_final.rows[0].cells[0]
set_cell_bg(fc, CDN)
pf = fc.paragraphs[0]
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_before = Pt(14)
pf.paragraph_format.space_after = Pt(6)
styled_run(pf, PAGE_28["citation"].replace('\n', ' '), bold=True, size_pt=10, color=CWH)
pf2 = fc.add_paragraph()
pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf2.paragraph_format.space_before = Pt(4)
pf2.paragraph_format.space_after = Pt(14)
styled_run(pf2, PAGE_28["pied"], size_pt=8, color='7799BB')

# ══════════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f'DOCX genere : {OUT}')
print(f'Taille : {os.path.getsize(OUT)//1024} Ko')
