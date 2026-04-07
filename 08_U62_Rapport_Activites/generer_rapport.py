#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur du rapport U62 - BAHAFID Mohamed - BTS MEC Session 2026
Convertit le fichier .md source en document .docx professionnel via python-docx.
"""

import os
import re
import sys
from pathlib import Path

# Fix Windows console encoding
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === CONFIGURATION ===
BASE_DIR = Path(__file__).parent
MD_FILE = BASE_DIR / "Rapport_Redaction" / "RAPPORT_U62_BAHAFID_Mohamed.md"
OUTPUT_FILE = BASE_DIR / "Rapport_Redaction" / "RAPPORT_U62_BAHAFID_Mohamed.docx"
IMAGES_DIR = BASE_DIR / "Annexes" / "Images"
PHOTOS_DIR = BASE_DIR / "Documents_Entreprise" / "Photos_Chantier"

# Couleurs
BLEU_FONCE = RGBColor(0x1B, 0x3A, 0x5C)
BLEU_MOYEN = RGBColor(0x2C, 0x5F, 0x8A)
BLEU_CLAIR = RGBColor(0xD6, 0xE8, 0xF7)
GRIS_FONCE = RGBColor(0x33, 0x33, 0x33)
GRIS_CLAIR = RGBColor(0xF2, 0xF2, 0xF2)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
NOIR = RGBColor(0x00, 0x00, 0x00)

# Photos de chantier sélectionnées (représentatives des projets)
PHOTOS_CHANTIER = [
    "20180228_143328.jpg",     # Travaux de terrassement
    "20180330_153115.jpg",     # Travaux de voirie
    "20180503_105806.jpg",     # Chantier d'assainissement
    "20180523_102115.jpg",     # Travaux de chaussée
    "20180604_133353.jpg",     # Travaux de trottoirs
    "20181206_124129.jpg",     # Ouvrages en béton
    "20181206_130504.jpg",     # Éclairage public
    "20181206_141917.jpg",     # Vue d'ensemble chantier
]


def setup_styles(doc):
    """Configure les styles du document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)
    font.color.rgb = GRIS_FONCE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = BLEU_FONCE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = BLEU_MOYEN
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = BLEU_MOYEN
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # Heading 4 (for sub-sub-sections like a), b), c)...)
    h4 = doc.styles["Heading 4"]
    h4.font.name = "Arial"
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = BLEU_MOYEN
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)
    h4.paragraph_format.keep_with_next = True


def setup_page(doc):
    """Configure la mise en page du document."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def add_header_footer(doc):
    """Ajoute en-tête et pied de page."""
    for section in doc.sections:
        # En-tête
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("Rapport U62 – BAHAFID Mohamed – BTS MEC Session 2026")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = "Arial"
        run.font.italic = True

        # Pied de page avec numéro de page
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Trait de séparation (via bordure)
        run = fp.add_run("BAHAFID Mohamed – Candidat n° 02537399911 – Académie de Lyon")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = "Arial"

        # Ajouter tabulation + numéro de page
        run2 = fp.add_run("    |    Page ")
        run2.font.size = Pt(7)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run2.font.name = "Arial"

        # Champ numéro de page
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fp.runs[-1]._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fp.runs[-1]._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        fp.runs[-1]._r.append(fldChar2)


def add_page_de_garde(doc):
    """Crée la page de garde professionnelle."""
    # Espace supérieur
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Logo BIMCO (si disponible)
    logo_path = IMAGES_DIR / "logo_bimco.png"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_path), width=Cm(4))
        p.paragraph_format.space_after = Pt(20)

    # Ligne décorative
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = BLEU_FONCE
    run.font.size = Pt(14)

    # Titre principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("RAPPORT U62")
    run.font.name = "Arial"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BLEU_FONCE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("COMPTE RENDU D'ACTIVITÉS PROFESSIONNELLES")
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLEU_MOYEN

    # Sous-titre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("BTS Management Économique de la Construction")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = GRIS_FONCE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("Session 2026")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BLEU_MOYEN

    # Ligne décorative
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = BLEU_FONCE
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(20)

    # Photo du candidat (si disponible)
    photo_path = IMAGES_DIR / "photo_bahafid.jpg"
    if photo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(photo_path), width=Cm(3.5))
        p.paragraph_format.space_after = Pt(15)

    # Infos candidat
    infos = [
        ("Candidat", "BAHAFID Mohamed"),
        ("N° Candidat", "02537399911"),
        ("Académie", "Lyon"),
        ("", ""),
        ("Structure d'accueil", "Conseil Régional de Béni Mellal-Khénifra (Maroc)"),
        ("Direction", "Agence d'Exécution des Projets"),
        ("Poste", "Technicien Études et Suivi des Travaux"),
        ("Durée d'expérience", "8 ans dans le BTP (3 ans Maroc + 5 ans France)"),
        ("", ""),
        ("Activité actuelle", "BIMCO – Projeteur BIM / Économiste de la construction"),
        ("SIREN", "999580053  |  Code APE : 7112B"),
    ]

    table = doc.add_table(rows=len([i for i in infos if i[0]]), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_idx = 0
    for label, value in infos:
        if not label:
            continue
        row = table.rows[row_idx]
        # Label cell
        cell_l = row.cells[0]
        cell_l.width = Cm(5)
        pl = cell_l.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rl = pl.add_run(label + " : ")
        rl.font.name = "Arial"
        rl.font.size = Pt(10)
        rl.font.bold = True
        rl.font.color.rgb = BLEU_FONCE
        # Value cell
        cell_v = row.cells[1]
        cell_v.width = Cm(10)
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(value)
        rv.font.name = "Arial"
        rv.font.size = Pt(10)
        rv.font.color.rgb = GRIS_FONCE
        row_idx += 1

    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Pied de page de garde
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.color.rgb = BLEU_FONCE
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Compte rendu d'activités professionnelles")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GRIS_FONCE

    # Saut de page
    doc.add_page_break()


def add_sommaire(doc):
    """Ajoute le sommaire."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("SOMMAIRE")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLEU_FONCE

    sommaire_items = [
        (0, "INTRODUCTION"),
        (0, "PARTIE 1 : CADRE PROFESSIONNEL"),
        (1, "1.1 Mon parcours professionnel"),
        (1, "1.2 Le Conseil Régional de Béni Mellal-Khénifra"),
        (1, "1.3 BIMCO – Mon activité indépendante"),
        (0, "PARTIE 2 : ACTIVITÉS ET PROJETS RÉALISÉS"),
        (1, "2.1 Projet 1 : Travaux de mise à niveau des centres des communes"),
        (2, "2.1.1 Présentation du projet"),
        (2, "2.1.2 Mes missions sur ce projet"),
        (2, "2.1.3 Aspects économiques du projet"),
        (2, "2.1.4 Difficultés rencontrées et solutions apportées"),
        (1, "2.2 Projet 2 : Construction de la RNC Lehri-Kerrouchen"),
        (2, "2.2.1 Présentation du projet"),
        (2, "2.2.2 Mes missions sur ce projet"),
        (2, "2.2.3 Difficultés rencontrées et solutions apportées"),
        (1, "2.3 Activités complémentaires"),
        (0, "PARTIE 3 : ANALYSE ET COMPÉTENCES"),
        (1, "3.1 Compétences acquises au regard du BTS MEC"),
        (1, "3.2 Analyse comparative Maroc / France"),
        (1, "3.3 Projet professionnel"),
        (0, "CONCLUSION"),
        (0, "ANNEXES"),
    ]

    for level, text in sommaire_items:
        p = doc.add_paragraph()
        indent = level * 1.0
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(3) if level > 0 else Pt(8)

        if level == 0:
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = BLEU_FONCE
        elif level == 1:
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = GRIS_FONCE
        else:
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Note pour l'utilisateur
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(Les numéros de page seront ajustés après mise en page finale dans Word)")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond à une cellule."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_from_data(doc, headers, rows, col_widths=None):
    """Crée un tableau formaté dans le document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style de bordure pour tout le tableau
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="2C5F8A"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="2C5F8A"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="2C5F8A"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="2C5F8A"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0C4DE"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # En-têtes
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "1B3A5C")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.replace("**", ""))
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = BLANC

    # Données
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            # Alternance de couleur
            if i % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
            p = cell.paragraphs[0]
            clean_text = cell_text.replace("**", "")

            # Bold si le texte contenait **
            is_bold = "**" in cell_text
            run = p.add_run(clean_text)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = GRIS_FONCE
            if is_bold:
                run.font.bold = True

            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)

    # Appliquer les largeurs de colonnes si spécifiées
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def parse_md_table(lines, start_idx):
    """Parse un tableau markdown et retourne (headers, rows, end_idx)."""
    headers = []
    rows = []

    # Ligne d'en-tête
    header_line = lines[start_idx].strip()
    if "|" in header_line:
        headers = [c.strip() for c in header_line.split("|") if c.strip()]

    # Ligne séparateur (---|---)
    sep_idx = start_idx + 1
    if sep_idx < len(lines) and re.match(r"\s*\|[\s\-:|]+\|", lines[sep_idx]):
        pass  # skip separator
    else:
        return headers, rows, start_idx + 1

    # Lignes de données
    idx = sep_idx + 1
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith("|"):
            break
        cols = [c.strip() for c in line.split("|") if c.strip() != ""]
        # Handle empty cells from split
        raw = line.split("|")
        cols = [c.strip() for c in raw[1:-1]]  # skip first and last empty from split
        rows.append(cols)
        idx += 1

    return headers, rows, idx


def add_formatted_paragraph(doc, text, style="Normal"):
    """Ajoute un paragraphe avec gras/italique inline."""
    p = doc.add_paragraph(style=style)
    # Parse bold and italic markers
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = p.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p


def add_image_safe(doc, image_path, width_cm=14, caption=None):
    """Insère une image si elle existe, avec légende optionnelle."""
    path = Path(image_path)
    if not path.exists():
        print(f"  [!] Image non trouvée : {path}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Image non disponible : {path.name}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(9)
        run.font.italic = True
        return

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        print(f"  [OK] Image insérée : {path.name}")
    except Exception as e:
        print(f"  [!] Erreur insertion image {path.name}: {e}")
        p = doc.add_paragraph()
        run = p.add_run(f"[Erreur image : {path.name}]")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = pc.add_run(caption)
        rc.font.name = "Arial"
        rc.font.size = Pt(9)
        rc.font.italic = True
        rc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        pc.paragraph_format.space_after = Pt(10)


def process_md_to_docx(doc, md_content):
    """Convertit le contenu markdown en éléments docx."""
    lines = md_content.split("\n")
    i = 0
    skip_header = True  # Skip the YAML-like header section until first ##

    # Track what we've already added (page de garde, sommaire)
    skip_sommaire = True
    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Sauter les lignes vides au début et la section titre/sommaire
        # (déjà gérés par page de garde et sommaire)
        if skip_header:
            if stripped.startswith("## INTRODUCTION"):
                skip_header = False
                # Don't skip this line, process it
            else:
                i += 1
                continue

        # Code blocks (organigramme, etc.)
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                # Add code block as formatted text
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    code_text = "\n".join(code_lines)
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(8)
                    run.font.color.rgb = GRIS_FONCE
                    # Grey background
                    shading_elm = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
                    )
                    p._element.get_or_add_pPr().append(shading_elm)
                in_code_block = False
                i += 1
                continue

        # Lignes séparateur ---
        if re.match(r"^---+\s*$", stripped):
            i += 1
            continue

        # Lignes vides
        if not stripped:
            i += 1
            continue

        # Images markdown ![alt](path)
        img_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            # Résoudre le chemin relatif
            full_path = (MD_FILE.parent / img_path).resolve()
            add_image_safe(doc, full_path, width_cm=12, caption=alt_text)
            i += 1
            continue

        # Placeholders [À FOURNIR...]
        if stripped.startswith("[À FOURNIR"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(stripped)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            # Background orange clair
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="FFF3E0" w:val="clear"/>'
            )
            p._element.get_or_add_pPr().append(shading_elm)
            i += 1
            continue

        # Headings ## = H1, ### = H2, #### = H3
        h_match = re.match(r"^(#{2,5})\s+(.*)", stripped)
        if h_match:
            level = len(h_match.group(1))
            title = h_match.group(2).strip()

            # ## → Heading 1, ### → Heading 2, #### → Heading 3
            if level == 2:
                # Parties principales
                heading_style = "Heading 1"
                # Don't page-break for INTRODUCTION (first heading)
                h = doc.add_heading(title, level=1)
                if title == "INTRODUCTION":
                    h.paragraph_format.page_break_before = False
            elif level == 3:
                doc.add_heading(title, level=2)
            elif level == 4:
                doc.add_heading(title, level=3)
            elif level == 5:
                doc.add_heading(title, level=4)
            i += 1
            continue

        # Tableaux markdown
        if "|" in stripped and stripped.startswith("|"):
            headers, rows, end_idx = parse_md_table(lines, i)
            if headers and rows:
                add_table_from_data(doc, headers, rows)
                doc.add_paragraph()  # Espace après tableau
            i = end_idx
            continue

        # Listes à puces
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = add_formatted_paragraph(doc, text)
            p.style = doc.styles["List Bullet"]
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Listes numérotées
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            text = num_match.group(2)
            p = add_formatted_paragraph(doc, text)
            p.style = doc.styles["List Number"]
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Paragraphes normaux (avec gras/italique)
        add_formatted_paragraph(doc, stripped)
        i += 1


def add_photos_chantier(doc):
    """Ajoute une sélection de photos de chantier en annexe."""
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("Annexe 9 – Sélection de photos de chantier")
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLEU_FONCE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Projets suivis au Conseil Régional de Béni Mellal-Khénifra (2017-2021)")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GRIS_FONCE
    p.paragraph_format.space_after = Pt(20)

    captions = [
        "Travaux de terrassement – Préparation de la plateforme",
        "Travaux de voirie – Mise en œuvre des couches de chaussée",
        "Chantier d'assainissement – Pose de canalisations",
        "Travaux de chaussée – Couche de fondation",
        "Aménagement des trottoirs – Pose de bordures",
        "Ouvrages en béton armé – Coffrage et bétonnage",
        "Éclairage public – Installation des candélabres",
        "Vue d'ensemble d'un chantier d'aménagement urbain",
    ]

    inserted = 0
    for idx, photo_name in enumerate(PHOTOS_CHANTIER):
        photo_path = PHOTOS_DIR / photo_name
        caption = captions[idx] if idx < len(captions) else f"Photo de chantier {idx+1}"
        if photo_path.exists():
            add_image_safe(doc, photo_path, width_cm=12, caption=caption)
            inserted += 1
        else:
            print(f"  [!] Photo de chantier non trouvée : {photo_name}")

    if inserted == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Photos de chantier à insérer manuellement]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    print(f"  Photos de chantier insérées : {inserted}/{len(PHOTOS_CHANTIER)}")


def main():
    """Point d'entrée principal."""
    print("=" * 60)
    print("  GÉNÉRATION DU RAPPORT U62 – BAHAFID Mohamed")
    print("  BTS MEC – Session 2026")
    print("=" * 60)

    # Vérifier que le fichier source existe
    if not MD_FILE.exists():
        print(f"\n[ERREUR] Fichier source introuvable : {MD_FILE}")
        return

    print(f"\n[1/6] Lecture du fichier source : {MD_FILE.name}")
    md_content = MD_FILE.read_text(encoding="utf-8")
    print(f"  → {len(md_content)} caractères, {md_content.count(chr(10))} lignes")

    # Créer le document
    print("\n[2/6] Création du document Word...")
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    print("\n[3/6] Génération de la page de garde...")
    add_page_de_garde(doc)

    print("\n[4/6] Génération du sommaire...")
    add_sommaire(doc)

    print("\n[5/6] Conversion du contenu markdown...")
    # Ajouter en-têtes/pieds de page
    add_header_footer(doc)
    process_md_to_docx(doc, md_content)

    print("\n[6/6] Ajout des photos de chantier en annexe...")
    add_photos_chantier(doc)

    # Sauvegarder
    print(f"\n{'='*60}")
    print(f"  Sauvegarde : {OUTPUT_FILE}")
    doc.save(str(OUTPUT_FILE))

    # Statistiques
    nb_paragraphs = len(doc.paragraphs)
    nb_tables = len(doc.tables)
    file_size = OUTPUT_FILE.stat().st_size / 1024

    print(f"\n  Statistiques du document :")
    print(f"  → Paragraphes : {nb_paragraphs}")
    print(f"  → Tableaux    : {nb_tables}")
    print(f"  → Taille      : {file_size:.0f} Ko")
    print(f"\n{'='*60}")
    print("  RAPPORT GÉNÉRÉ AVEC SUCCÈS !")
    print(f"{'='*60}")

    print("\n  À FAIRE par le candidat :")
    print("  [ ] Prendre 3-4 captures d'écran de l'appli Gestion Chantiers")
    print("  [ ] Capture Google Maps des 4 communes de Khénifra")
    print("  [ ] Capture Google Maps de la route Lehri-Kerrouchen")
    print("  [ ] Ajouter les attestations de travail en annexe")
    print("  [ ] Ajouter l'attestation de formation BIM en annexe")
    print("  [ ] Relecture finale + pagination du sommaire dans Word")
    print("  [ ] Impression et reliure")


if __name__ == "__main__":
    main()
