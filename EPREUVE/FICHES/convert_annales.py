import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches

files = [
    'ANNALE_E4_2025_CORRIGE',
    'ANNALE_E52_2024_CORRIGE',
    'FICHE_MEMO_A4',
    'PLANNING_REVISION',
]

base = 'D:/PREPA BTS MEC/EPREUVE/FICHES'

for name in files:
    md_path = os.path.join(base, name + '.md')
    docx_path = os.path.join(base, name + '.docx')

    doc = Document()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_rows = []
    in_code = False

    def flush_table():
        global in_table, table_rows
        if not table_rows:
            in_table = False
            return
        data_rows = [r for r in table_rows if not all(c.strip().replace('-','').replace('|','') == '' for c in r)]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        max_cols = max(len(r) for r in data_rows)
        table = doc.add_table(rows=len(data_rows), cols=max_cols)
        table.style = 'Table Grid'
        for ri, row in enumerate(data_rows):
            for ci in range(max_cols):
                cell_text = row[ci].strip() if ci < len(row) else ''
                cell_obj = table.rows[ri].cells[ci]
                cell_obj.text = cell_text
                if ri == 0:
                    for para in cell_obj.paragraphs:
                        for run in para.runs:
                            run.bold = True
        in_table = False
        table_rows = []

    for line in lines:
        line_raw = line.rstrip()

        # Code blocks
        if line_raw.strip().startswith('```'):
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph(line_raw)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        # Tables
        if '|' in line_raw and line_raw.strip().startswith('|'):
            cells = line_raw.strip().split('|')
            cells = [c for c in cells if c != '']
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        else:
            if in_table:
                flush_table()

        stripped = line_raw.strip()

        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.italic = True
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('---'):
            doc.add_paragraph('=' * 50)
        elif stripped == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(stripped)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f'OK -> {name}.docx')

print('Conversion terminee.')
