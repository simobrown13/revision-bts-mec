"""
Conversion des fiches Markdown vers Word (.docx)
Utilise python-docx pour générer des fichiers Word formatés.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FICHES_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Couleurs thématiques ──────────────────────────────────────────────────────
COLOR_H1      = RGBColor(0x1F, 0x49, 0x7D)   # bleu foncé
COLOR_H2      = RGBColor(0x2E, 0x74, 0xB5)   # bleu moyen
COLOR_H3      = RGBColor(0x5B, 0x9B, 0xD5)   # bleu clair
COLOR_CODE    = RGBColor(0x2B, 0x57, 0x2A)   # vert foncé
COLOR_CODE_BG = RGBColor(0xF0, 0xF4, 0xF0)   # fond vert très clair
COLOR_TABLE_H = RGBColor(0x1F, 0x49, 0x7D)   # en-tête table = bleu foncé


def set_cell_bg(cell, color_hex: str):
    """Applique une couleur de fond à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_borders(table):
    """Ajoute des bordures légères à toutes les cellules."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "BFBFBF")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_run_with_inline(para, text: str, base_size: int = 11, bold_default=False):
    """
    Ajoute du texte dans un paragraphe en gérant **gras**, *italique*,
    `code inline` et les émojis ⭐✓✗.
    """
    # Découpages par marqueurs inline
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(base_size)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = COLOR_CODE
        else:
            run = para.add_run(part)
            run.bold = bold_default
            run.font.size = Pt(base_size)


def parse_and_add_table(doc: Document, lines: list, start: int):
    """
    Détecte et insère un tableau Markdown dans le document Word.
    Retourne l'index de la dernière ligne consommée.
    """
    # Collecter les lignes du tableau
    table_lines = []
    i = start
    while i < len(lines) and lines[i].startswith("|"):
        table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 2:
        return start  # pas un vrai tableau

    # Parser l'en-tête
    header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]
    # Ligne séparateur (index 1) → ignorer
    data_rows = []
    for row_line in table_lines[2:]:
        cells = [c.strip() for c in row_line.split("|") if c.strip() != ""]
        if cells:
            data_rows.append(cells)

    ncols = len(header_cells)
    nrows = 1 + len(data_rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # En-tête
    for j, cell_text in enumerate(header_cells):
        cell = table.cell(0, j)
        set_cell_bg(cell, "1F497D")
        para = cell.paragraphs[0]
        run = para.add_run(re.sub(r'\*\*|\*|`', '', cell_text))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Données
    for r, row_data in enumerate(data_rows):
        bg = "EAF0F8" if r % 2 == 0 else "FFFFFF"
        for j in range(ncols):
            cell = table.cell(r + 1, j)
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            text = row_data[j] if j < len(row_data) else ""
            add_run_with_inline(para, text, base_size=10)

    set_cell_borders(table)
    doc.add_paragraph()  # espace après tableau
    return i - 1


def convert_md_to_docx(md_path: str, docx_path: str):
    doc = Document()

    # ── Marges du document ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    with open(md_path, encoding="utf-8") as f:
        raw = f.read()

    lines = raw.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Titre H1 ──────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_H1
            # Ligne de séparation sous le titre
            doc.add_paragraph()

        # ── Titre H2 ──────────────────────────────────────────────────────
        elif line.startswith("## "):
            text = line[3:].strip()
            para = doc.add_heading(level=2)
            para.clear()
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_H2

        # ── Titre H3 ──────────────────────────────────────────────────────
        elif line.startswith("### "):
            text = line[4:].strip()
            para = doc.add_heading(level=3)
            para.clear()
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_H3

        # ── Titre H4 ──────────────────────────────────────────────────────
        elif line.startswith("#### "):
            text = line[5:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            run.underline = True
            run.font.size = Pt(11)

        # ── Ligne de séparation ───────────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pb = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "2E74B5")
            pb.append(bottom)
            pPr.append(pb)

        # ── Tableau Markdown ──────────────────────────────────────────────
        elif line.startswith("|"):
            i = parse_and_add_table(doc, lines, i)

        # ── Bloc de code ──────────────────────────────────────────────────
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            if code_text.strip():
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_CODE
                # Fond léger sur le paragraphe
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F0F4F0")
                pPr.append(shd)
                para.paragraph_format.left_indent = Cm(0.5)
            doc.add_paragraph()

        # ── Liste à puces (- ou * en début de ligne) ──────────────────────
        elif re.match(r'^[-*] ', line):
            text = line[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            add_run_with_inline(para, text, base_size=11)

        # ── Sous-liste (  - ou    - ) ─────────────────────────────────────
        elif re.match(r'^ {2,}[-*] ', line):
            text = re.sub(r'^ +[-*] ', '', line)
            para = doc.add_paragraph(style="List Bullet 2")
            add_run_with_inline(para, text, base_size=10)

        # ── Liste numérotée ───────────────────────────────────────────────
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            para = doc.add_paragraph(style="List Number")
            add_run_with_inline(para, text, base_size=11)

        # ── Citation (>) ──────────────────────────────────────────────────
        elif line.startswith("> "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        # ── Ligne vide ────────────────────────────────────────────────────
        elif line.strip() == "":
            # Ajouter un espace seulement si le dernier paragraphe n'est pas vide
            if doc.paragraphs and doc.paragraphs[-1].text.strip() != "":
                doc.add_paragraph()

        # ── Paragraphe normal ────────────────────────────────────────────
        else:
            if line.strip():
                para = doc.add_paragraph()
                add_run_with_inline(para, line.strip(), base_size=11)

        i += 1

    doc.save(docx_path)
    print(f"  OK  {os.path.basename(md_path):45s} -> {os.path.basename(docx_path)}")


def main():
    md_files = [f for f in os.listdir(FICHES_DIR) if f.endswith(".md")]
    md_files.sort()

    print(f"\nConversion de {len(md_files)} fichiers Markdown -> Word\n")
    print("=" * 60)

    for md_file in md_files:
        md_path   = os.path.join(FICHES_DIR, md_file)
        docx_name = md_file.replace(".md", ".docx")
        docx_path = os.path.join(FICHES_DIR, docx_name)
        convert_md_to_docx(md_path, docx_path)

    print("=" * 60)
    print(f"\nTermine ! {len(md_files)} fichiers Word crees dans :\n{FICHES_DIR}\n")


if __name__ == "__main__":
    main()
