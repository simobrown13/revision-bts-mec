"""
Conversion speciale de FORMULES_ESSENTIELLES.md en Word
avec mise en page soignee : encadres colores par section.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FICHES_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE  = os.path.join(FICHES_DIR, "FORMULES_ESSENTIELLES.md")
OUT_FILE = os.path.join(FICHES_DIR, "FORMULES_ESSENTIELLES.docx")

# Palette de couleurs par section
SECTION_COLORS = {
    "THERMIQUE":  ("1F497D", "D6E4F0"),   # bleu
    "ACOUSTIQUE": ("375623", "E2EFDA"),   # vert
    "OFFRES":     ("7B3F00", "FCE8D5"),   # orange
    "PRIX":       ("5C0070", "EFD9F5"),   # violet
    "AVANCE":     ("C55A11", "FFF2CC"),   # jaune-orange
    "RETENUE":    ("C00000", "FFDEDE"),   # rouge
    "DECOMPTE":   ("1F497D", "D6E4F0"),   # bleu
    "RATIOS":     ("375623", "E2EFDA"),   # vert
    "RENTAB":     ("7B3F00", "FCE8D5"),   # orange
    "RECAPITU":   ("2F4F4F", "F0F0F0"),   # gris
}

DEFAULT_COLORS = ("2E74B5", "EAF0F8")


def get_section_colors(title: str):
    t = title.upper()
    for key, colors in SECTION_COLORS.items():
        if key in t:
            return colors
    return DEFAULT_COLORS


def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_borders(table, color="BFBFBF"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color)
                tcBorders.append(b)
            tcPr.append(tcBorders)


def add_inline(para, text: str, size=11, bold=False, color=None, mono=False):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    for part in pattern.split(text):
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = para.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.color.rgb = RGBColor(0x2B, 0x57, 0x2A)
        else:
            r = para.add_run(part)
        r.bold = bold if not r.bold else True
        r.font.size = Pt(size)
        if mono:
            r.font.name = "Courier New"
        if color:
            r.font.color.rgb = hex_to_rgb(color)


def add_code_block(doc, code_text: str, bg_hex="F0F4F0", fg_hex="2B572A"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    cell.width = Inches(6.5)
    para = cell.paragraphs[0]
    para.paragraph_format.left_indent  = Cm(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(code_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = hex_to_rgb(fg_hex)
    # Bordure coloree a gauche (effet highlight)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), fg_hex)
    tcBorders.append(left)
    for side in ("top", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_header(doc, title: str, level=2):
    hdr_hex, _ = get_section_colors(title)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, hdr_hex)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.3)
    prefix = "E4 — " if "E4" in title else ("E52 — " if "E52" in title else "")
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(13 if level == 2 else 11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_formula_box(doc, formula_name: str, lines_code: list, lines_vars: list, section_title: str):
    """Insere un encadre formule : titre + code + variables."""
    hdr_hex, bg_hex = get_section_colors(section_title)

    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Ligne titre de la formule
    hdr_cell = table.cell(0, 0)
    set_cell_bg(hdr_cell, hdr_hex)
    hdr_p = hdr_cell.paragraphs[0]
    hdr_p.paragraph_format.left_indent = Cm(0.3)
    hdr_r = hdr_p.add_run(formula_name)
    hdr_r.bold = True
    hdr_r.font.size = Pt(10)
    hdr_r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Ligne contenu (code + variables)
    body_cell = table.cell(1, 0)
    set_cell_bg(body_cell, bg_hex)
    body_cell._tc.get_or_add_tcPr()

    for line in lines_code:
        p = body_cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.name  = "Courier New"
        r.font.size  = Pt(10)
        r.bold = True
        r.font.color.rgb = hex_to_rgb(hdr_hex)

    if lines_vars:
        sep = body_cell.add_paragraph()
        sep.paragraph_format.space_after = Pt(1)
        for var_line in lines_vars:
            if var_line.strip():
                p = body_cell.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(1)
                add_inline(p, var_line.strip(), size=9, color="404040")

    # Retirer le premier paragraphe vide cree par defaut dans la cellule
    if body_cell.paragraphs[0].text == "":
        p_el = body_cell.paragraphs[0]._p
        p_el.getparent().remove(p_el)

    set_table_borders(table, color="DDDDDD")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def parse_and_render(doc: Document, lines: list):
    i = 0
    current_section = ""

    while i < len(lines):
        line = lines[i]

        # H1 - Titre principal
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(text)
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            doc.add_paragraph()
            i += 1; continue

        # Citation > (sous-titre)
        if line.startswith("> "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(text)
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            doc.add_paragraph()
            i += 1; continue

        # H2 - Section principale
        if line.startswith("## "):
            current_section = line[3:].strip()
            add_section_header(doc, current_section, level=2)
            i += 1; continue

        # H3 - Sous-section = une formule
        if line.startswith("### "):
            formula_name = line[4:].strip()
            # Collecter le contenu jusqu'au prochain H2/H3/---
            i += 1
            code_lines = []
            var_lines  = []
            in_code    = False

            while i < len(lines):
                l = lines[i]
                if l.startswith("## ") or l.startswith("### ") or l.startswith("---"):
                    break
                if l.startswith("```"):
                    in_code = not in_code
                    i += 1; continue
                if in_code:
                    code_lines.append(l)
                elif l.strip().startswith(("e =", "a =", "R =", "U =", "λ", "ρ", "Cp", "φ", "TR", "V =", "A =",
                                           "Si", "αi", "Im", "I0", "Cn", "P0", "Pr", "DS", "CD", "MB", "MBH",
                                           "RB", "FOp", "FG", "RG", "N =", "Exigence", "Cn ", "m ", "Le ",
                                           "Il se", "R du", "RG m", "Plaf", "Marg", "Mar", "Res", "Net",
                                           "TVA", "Prix_", "Déla", "Dela")):
                    var_lines.append(l)
                elif l.strip():
                    var_lines.append(l)
                i += 1

            add_formula_box(doc, formula_name, code_lines, var_lines, current_section)
            continue

        # Séparateur ---
        if line.strip() in ("---", "***"):
            doc.add_paragraph()
            i += 1; continue

        # Tableau
        if line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            if len(tbl_lines) >= 2:
                headers = [c.strip() for c in tbl_lines[0].split("|") if c.strip()]
                data_rows = []
                for rl in tbl_lines[2:]:
                    cells = [c.strip() for c in rl.split("|")]
                    cells = [c for c in cells if c != ""]
                    if cells:
                        data_rows.append(cells)
                ncols = len(headers)
                nrows = 1 + len(data_rows)
                table = doc.add_table(rows=nrows, cols=ncols)
                table.style = "Table Grid"
                hdr_hex, bg_hex = get_section_colors(current_section)
                for j, h in enumerate(headers):
                    cell = table.cell(0, j)
                    set_cell_bg(cell, hdr_hex)
                    p = cell.paragraphs[0]
                    r = p.add_run(re.sub(r'\*\*|\*|`', '', h))
                    r.bold = True; r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for ri, row_data in enumerate(data_rows):
                    bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
                    for j in range(ncols):
                        cell = table.cell(ri+1, j)
                        set_cell_bg(cell, bg)
                        p = cell.paragraphs[0]
                        txt = row_data[j] if j < len(row_data) else ""
                        add_inline(p, txt, size=9)
                set_table_borders(table)
                doc.add_paragraph()
            continue

        # Ligne vide
        if not line.strip():
            i += 1; continue

        # Texte normal
        para = doc.add_paragraph()
        add_inline(para, line.strip(), size=11)
        i += 1


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    with open(MD_FILE, encoding="utf-8") as f:
        lines = f.read().splitlines()

    parse_and_render(doc, lines)
    doc.save(OUT_FILE)
    print(f"Fichier cree : {OUT_FILE}")


if __name__ == "__main__":
    main()
